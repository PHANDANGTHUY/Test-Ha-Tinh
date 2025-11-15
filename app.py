import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
import google.generativeai as genai
from io import BytesIO
import re
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

# Cấu hình trang
st.set_page_config(
    page_title="Hệ thống Thẩm định Phương án Kinh doanh",
    page_icon="🏦",
    layout="wide"
)

# CSS tùy chỉnh
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #e3f2fd 0%, #bbdefb 100%);
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .sub-header {
        font-size: 1.3rem;
        font-weight: 600;
        color: #0d47a1;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
        border-left: 4px solid #1f77b4;
        padding-left: 10px;
    }
    .metric-card {
        background-color: #f5f5f5;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #ddd;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        background-color: #e3f2fd;
        border-radius: 5px 5px 0 0;
        padding: 10px 20px;
        font-weight: 600;
    }
    .stTabs [aria-selected="true"] {
        background-color: #1f77b4;
        color: white;
    }
    .api-note {
        background-color: #fff3cd;
        border: 1px solid #ffc107;
        border-radius: 5px;
        padding: 10px;
        margin-bottom: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# Khởi tạo session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'data_changed' not in st.session_state:
    st.session_state.data_changed = False
if 'uploaded_content' not in st.session_state:
    st.session_state.uploaded_content = ""

# Hàm định dạng số
def format_number(num):
    """Định dạng số với dấu phẩy phân cách hàng nghìn"""
    if num == 0:
        return "0"
    return f"{int(num):,}"

def parse_number(text):
    """Chuyển đổi text có dấu phẩy thành số"""
    if isinstance(text, (int, float)):
        return float(text)
    return float(str(text).replace(",", ""))

# Hàm trích xuất thông tin từ file docx
def extract_info_from_docx(file):
    """Trích xuất thông tin từ file phương án kinh doanh"""
    doc = Document(file)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Lưu nội dung gốc vào session state
    st.session_state.uploaded_content = full_text
    
    info = {
        'ho_ten': '',
        'cccd': '',
        'dia_chi': '',
        'sdt': '',
        'muc_dich_vay': '',
        'tong_nhu_cau_von': 0,
        'von_doi_ung': 0,
        'so_tien_vay': 0,
        'lai_suat': 0,
        'thoi_gian_vay': 0,
        'mo_ta_tai_san': '',
        'gia_tri_tai_san': 0,
        'thu_nhap_thang': 0,
        'chi_phi_thang': 0
    }
    
    # Trích xuất họ tên
    ho_ten_match = re.search(r'Họ và tên:\s*([^\-\n]+)', full_text)
    if ho_ten_match:
        info['ho_ten'] = ho_ten_match.group(1).strip()
    
    # Trích xuất CCCD
    cccd_match = re.search(r'(?:CMND/CCCD|CCCD).*?:\s*(\d+)', full_text)
    if cccd_match:
        info['cccd'] = cccd_match.group(1).strip()
    
    # Trích xuất địa chỉ
    dia_chi_match = re.search(r'Nơi cư trú:\s*([^\n]+)', full_text)
    if dia_chi_match:
        info['dia_chi'] = dia_chi_match.group(1).strip()
    
    # Trích xuất SĐT
    sdt_match = re.search(r'Số điện thoại:\s*(\d+)', full_text)
    if sdt_match:
        info['sdt'] = sdt_match.group(1).strip()
    
    # Trích xuất mục đích vay
    muc_dich_match = re.search(r'Mục đích vay:\s*([^\n]+)', full_text)
    if muc_dich_match:
        info['muc_dich_vay'] = muc_dich_match.group(1).strip()
    
    # Trích xuất tổng nhu cầu vốn
    tong_von_match = re.search(r'Tổng nhu cầu vốn:\s*([\d.,]+)', full_text)
    if tong_von_match:
        info['tong_nhu_cau_von'] = float(tong_von_match.group(1).replace('.', '').replace(',', ''))
    
    # Trích xuất vốn đối ứng
    von_du_match = re.search(r'Vốn đối ứng.*?:\s*([\d.,]+)', full_text)
    if von_du_match:
        info['von_doi_ung'] = float(von_du_match.group(1).replace('.', '').replace(',', ''))
    
    # Trích xuất số tiền vay
    tien_vay_match = re.search(r'Vốn vay.*?số tiền:\s*([\d.,]+)', full_text)
    if tien_vay_match:
        info['so_tien_vay'] = float(tien_vay_match.group(1).replace('.', '').replace(',', ''))
    
    # Trích xuất lãi suất
    lai_suat_match = re.search(r'Lãi suất:\s*([\d.,]+)', full_text)
    if lai_suat_match:
        info['lai_suat'] = float(lai_suat_match.group(1).replace(',', '.'))
    
    # Trích xuất thời gian vay
    thoi_gian_match = re.search(r'Thời hạn vay:\s*(\d+)', full_text)
    if thoi_gian_match:
        info['thoi_gian_vay'] = int(thoi_gian_match.group(1))
    
    # Trích xuất giá trị tài sản
    tai_san_match = re.search(r'Giá trị:?\s*([\d.,]+)\s*đồng', full_text)
    if tai_san_match:
        info['gia_tri_tai_san'] = float(tai_san_match.group(1).replace('.', '').replace(',', ''))
    
    # Trích xuất thu nhập tháng
    thu_nhap_match = re.search(r'Tổng thu nhập.*?:\s*([\d.,]+)', full_text)
    if thu_nhap_match:
        info['thu_nhap_thang'] = float(thu_nhap_match.group(1).replace('.', '').replace(',', ''))
    
    # Trích xuất chi phí tháng
    chi_phi_match = re.search(r'Tổng chi phí.*?:\s*([\d.,]+)', full_text)
    if chi_phi_match:
        info['chi_phi_thang'] = float(chi_phi_match.group(1).replace('.', '').replace(',', ''))
    
    # Mô tả tài sản
    if 'Bất động sản' in full_text:
        info['mo_ta_tai_san'] = 'Bất động sản (nhà và đất)'
    
    return info

# Hàm tính toán kế hoạch trả nợ
def calculate_repayment_schedule(so_tien_vay, lai_suat_nam, thoi_han_thang):
    """Tính toán kế hoạch trả nợ theo phương thức trả gốc đều"""
    if so_tien_vay <= 0 or thoi_han_thang <= 0:
        return pd.DataFrame()
    
    lai_suat_thang = lai_suat_nam / 12 / 100
    goc_tra_moi_ky = so_tien_vay / thoi_han_thang
    
    schedule = []
    du_no = so_tien_vay
    
    for ky in range(1, thoi_han_thang + 1):
        lai_tra = du_no * lai_suat_thang
        tong_tra = goc_tra_moi_ky + lai_tra
        du_no_cuoi = du_no - goc_tra_moi_ky
        
        schedule.append({
            'Kỳ': ky,
            'Dư nợ đầu kỳ': format_number(du_no),
            'Gốc trả': format_number(goc_tra_moi_ky),
            'Lãi trả': format_number(lai_tra),
            'Tổng trả': format_number(tong_tra),
            'Dư nợ cuối kỳ': format_number(max(0, du_no_cuoi))
        })
        
        du_no = du_no_cuoi
    
    return pd.DataFrame(schedule)

# Hàm tạo file Excel kế hoạch trả nợ
def create_repayment_excel(df, customer_info):
    """Tạo file Excel kế hoạch trả nợ"""
    output = BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Kế hoạch trả nợ"
    
    # Tiêu đề
    ws['A1'] = "KẾ HOẠCH TRẢ NỢ"
    ws['A1'].font = Font(size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')
    ws.merge_cells('A1:F1')
    
    # Thông tin khách hàng
    ws['A3'] = f"Khách hàng: {customer_info['ho_ten']}"
    ws['A4'] = f"Số tiền vay: {customer_info['so_tien_vay']} VNĐ"
    ws['A5'] = f"Lãi suất: {customer_info['lai_suat']}%/năm"
    ws['A6'] = f"Thời hạn: {customer_info['thoi_gian_vay']} tháng"
    
    # Header bảng
    headers = df.columns.tolist()
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=8, column=col_num, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F77B4", end_color="1F77B4", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')
    
    # Dữ liệu
    for row_num, row_data in enumerate(df.values, 9):
        for col_num, value in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col_num, value=value)
    
    wb.save(output)
    output.seek(0)
    return output

# Hàm tạo báo cáo thẩm định
def create_assessment_report(data):
    """Tạo báo cáo thẩm định dưới dạng Word"""
    doc = Document()
    
    # Tiêu đề
    heading = doc.add_heading('BÁO CÁO THẨM ĐỊNH PHƯƠNG ÁN KINH DOANH', 0)
    heading.alignment = 1  # Center
    
    # Thông tin khách hàng
    doc.add_heading('I. THÔNG TIN KHÁCH HÀNG', 1)
    doc.add_paragraph(f"Họ và tên: {data['ho_ten']}")
    doc.add_paragraph(f"CCCD: {data['cccd']}")
    doc.add_paragraph(f"Địa chỉ: {data['dia_chi']}")
    doc.add_paragraph(f"Số điện thoại: {data['sdt']}")
    
    # Thông tin phương án vay
    doc.add_heading('II. THÔNG TIN PHƯƠNG ÁN VAY', 1)
    doc.add_paragraph(f"Mục đích vay: {data['muc_dich_vay']}")
    doc.add_paragraph(f"Tổng nhu cầu vốn: {format_number(data['tong_nhu_cau_von'])} VNĐ")
    doc.add_paragraph(f"Vốn đối ứng: {format_number(data['von_doi_ung'])} VNĐ")
    doc.add_paragraph(f"Số tiền vay: {format_number(data['so_tien_vay'])} VNĐ")
    doc.add_paragraph(f"Lãi suất: {data['lai_suat']}%/năm")
    doc.add_paragraph(f"Thời gian vay: {data['thoi_gian_vay']} tháng")
    
    # Phân tích tài chính
    doc.add_heading('III. PHÂN TÍCH TÀI CHÍNH', 1)
    ty_le_vay = (data['so_tien_vay'] / data['tong_nhu_cau_von'] * 100) if data['tong_nhu_cau_von'] > 0 else 0
    ty_le_doi_ung = (data['von_doi_ung'] / data['tong_nhu_cau_von'] * 100) if data['tong_nhu_cau_von'] > 0 else 0
    
    doc.add_paragraph(f"Tỷ lệ vay/Tổng nhu cầu vốn: {ty_le_vay:.2f}%")
    doc.add_paragraph(f"Tỷ lệ vốn đối ứng: {ty_le_doi_ung:.2f}%")
    
    if data['gia_tri_tai_san'] > 0:
        ltv = (data['so_tien_vay'] / data['gia_tri_tai_san'] * 100)
        doc.add_paragraph(f"LTV (Loan to Value): {ltv:.2f}%")
    
    # Khả năng trả nợ
    doc.add_heading('IV. ĐÁNH GIÁ KHẢ NĂNG TRẢ NỢ', 1)
    doc.add_paragraph(f"Thu nhập hàng tháng: {format_number(data['thu_nhap_thang'])} VNĐ")
    doc.add_paragraph(f"Chi phí hàng tháng: {format_number(data['chi_phi_thang'])} VNĐ")
    
    thu_nhap_rong = data['thu_nhap_thang'] - data['chi_phi_thang']
    doc.add_paragraph(f"Thu nhập ròng: {format_number(thu_nhap_rong)} VNĐ")
    
    # Tài sản đảm bảo
    doc.add_heading('V. TÀI SẢN ĐẢM BẢO', 1)
    doc.add_paragraph(f"Mô tả: {data['mo_ta_tai_san']}")
    doc.add_paragraph(f"Giá trị định giá: {format_number(data['gia_tri_tai_san'])} VNĐ")
    
    # Kết luận
    doc.add_heading('VI. KẾT LUẬN', 1)
    doc.add_paragraph(f"Ngày lập báo cáo: {datetime.now().strftime('%d/%m/%Y')}")
    
    # Lưu vào BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Header
st.markdown('<div class="main-header">🏦 HỆ THỐNG THẨM ĐỊNH PHƯƠNG ÁN KINH DOANH</div>', unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.markdown("### ⚙️ Cấu hình")
    
    # Hướng dẫn lấy API Key
    with st.expander("📖 Hướng dẫn lấy Gemini API Key (MIỄN PHÍ)", expanded=False):
        st.markdown("""
        **Các bước lấy API Key:**
        
        1. Truy cập: [https://aistudio.google.com/app/apikey](https://aistudio.google.com/app/apikey)
        2. Đăng nhập bằng tài khoản Google
        3. Nhấn nút **"Create API Key"**
        4. Chọn project hoặc tạo mới
        5. Copy API key và paste vào ô bên dưới
        
        ⚠️ **Lưu ý:**
        - API key hoàn toàn MIỄN PHÍ
        - Không chia sẻ API key với người khác
        - Nếu key hết hạn, tạo key mới
        """)
    
    # API Key input
    api_key = st.text_input(
        "🔑 Gemini API Key", 
        type="password", 
        help="Nhập API key của bạn để sử dụng tính năng AI",
        placeholder="AIza..."
    )
    
    if api_key:
        try:
            genai.configure(api_key=api_key)
            # Test API key
            model = genai.GenerativeModel('gemini-2.0-flash-exp')
            st.success("✅ API Key hợp lệ!")
        except Exception as e:
            if "API_KEY_INVALID" in str(e) or "expired" in str(e).lower():
                st.error("❌ API Key không hợp lệ hoặc đã hết hạn!")
                st.markdown("""
                <div class='api-note'>
                    <strong>💡 Giải pháp:</strong><br>
                    1. Kiểm tra lại API key đã copy đúng chưa<br>
                    2. Tạo API key mới tại: <a href='https://aistudio.google.com/app/apikey' target='_blank'>Google AI Studio</a><br>
                    3. Đảm bảo đã bật Gemini API trong project
                </div>
                """, unsafe_allow_html=True)
            else:
                st.error(f"❌ Lỗi: {str(e)}")
    else:
        st.info("ℹ️ Nhập API key để sử dụng tính năng AI")
    
    st.markdown("---")
    
    # Chức năng xuất dữ liệu
    st.markdown("### 📊 Xuất dữ liệu")
    export_option = st.selectbox(
        "Chọn loại xuất",
        ["-- Chọn --", "Xuất Kế hoạch trả nợ (Excel)", "Xuất Báo cáo Thẩm định"]
    )
    
    if export_option != "-- Chọn --":
        if st.button("Thực hiện", type="primary"):
            if 'customer_data' in st.session_state:
                data = st.session_state.customer_data
                
                if export_option == "Xuất Kế hoạch trả nợ (Excel)":
                    schedule = calculate_repayment_schedule(
                        data['so_tien_vay'],
                        data['lai_suat'],
                        data['thoi_gian_vay']
                    )
                    if not schedule.empty:
                        excel_file = create_repayment_excel(schedule, data)
                        st.download_button(
                            label="📥 Tải về Excel",
                            data=excel_file,
                            file_name=f"ke_hoach_tra_no_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                
                elif export_option == "Xuất Báo cáo Thẩm định":
                    report_file = create_assessment_report(data)
                    st.download_button(
                        label="📥 Tải về Báo cáo",
                        data=report_file,
                        file_name=f"bao_cao_tham_dinh_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("⚠️ Vui lòng nhập đầy đủ thông tin trước khi xuất dữ liệu!")

# Tabs chính
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📝 Nhập liệu & Trích xuất",
    "📊 Phân tích Chỉ số & Dòng tiền",
    "📈 Biểu đồ Trực quan",
    "🤖 Phân tích bởi AI",
    "💬 Chatbot Hỗ trợ"
])

# Tab 1: Nhập liệu & Trích xuất
with tab1:
    st.markdown('<div class="sub-header">📤 Upload File Phương án</div>', unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Chọn file phương án kinh doanh (.docx)",
        type=['docx'],
        help="Upload file phương án kinh doanh của khách hàng"
    )
    
    # Khởi tạo dữ liệu mặc định
    default_data = {
        'ho_ten': '',
        'cccd': '',
        'dia_chi': '',
        'sdt': '',
        'muc_dich_vay': '',
        'tong_nhu_cau_von': 0,
        'von_doi_ung': 0,
        'so_tien_vay': 0,
        'lai_suat': 0.0,
        'thoi_gian_vay': 0,
        'mo_ta_tai_san': '',
        'gia_tri_tai_san': 0,
        'thu_nhap_thang': 0,
        'chi_phi_thang': 0
    }
    
    # Trích xuất nếu có file upload
    if uploaded_file is not None:
        with st.spinner('🔄 Đang trích xuất thông tin từ file...'):
            extracted_data = extract_info_from_docx(uploaded_file)
            st.success('✅ Trích xuất thông tin thành công!')
    else:
        extracted_data = default_data
    
    st.markdown('<div class="sub-header">👤 Vùng 1 - Thông tin Khách hàng</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        ho_ten = st.text_input("Họ và tên", value=extracted_data['ho_ten'])
        cccd = st.text_input("CCCD", value=extracted_data['cccd'])
    with col2:
        dia_chi = st.text_area("Địa chỉ", value=extracted_data['dia_chi'], height=100)
        sdt = st.text_input("Số điện thoại", value=extracted_data['sdt'])
    
    st.markdown('<div class="sub-header">💰 Vùng 2 - Thông tin Phương án Vay</div>', unsafe_allow_html=True)
    
    muc_dich_vay = st.text_area("Mục đích vay", value=extracted_data['muc_dich_vay'], height=80)
    
    col3, col4, col5 = st.columns(3)
    with col3:
        tong_nhu_cau_von = st.number_input(
            "Tổng nhu cầu vốn (VNĐ)",
            min_value=0,
            value=int(extracted_data['tong_nhu_cau_von']),
            step=1000000,
            format="%d"
        )
        st.caption(f"💵 {format_number(tong_nhu_cau_von)} VNĐ")
    
    with col4:
        von_doi_ung = st.number_input(
            "Vốn đối ứng (VNĐ)",
            min_value=0,
            value=int(extracted_data['von_doi_ung']),
            step=1000000,
            format="%d"
        )
        st.caption(f"💵 {format_number(von_doi_ung)} VNĐ")
    
    with col5:
        so_tien_vay = st.number_input(
            "Số tiền vay (VNĐ)",
            min_value=0,
            value=int(extracted_data['so_tien_vay']),
            step=1000000,
            format="%d"
        )
        st.caption(f"💵 {format_number(so_tien_vay)} VNĐ")
    
    col6, col7 = st.columns(2)
    with col6:
        lai_suat = st.number_input(
            "Lãi suất (%/năm)",
            min_value=0.0,
            max_value=100.0,
            value=float(extracted_data['lai_suat']),
            step=0.1,
            format="%.2f"
        )
    
    with col7:
        thoi_gian_vay = st.number_input(
            "Thời gian vay (tháng)",
            min_value=0,
            max_value=360,
            value=int(extracted_data['thoi_gian_vay']),
            step=1
        )
    
    st.markdown('<div class="sub-header">🏠 Vùng 3 - Thông tin Tài sản Đảm bảo</div>', unsafe_allow_html=True)
    
    col8, col9 = st.columns(2)
    with col8:
        mo_ta_tai_san = st.text_area(
            "Mô tả tài sản",
            value=extracted_data['mo_ta_tai_san'],
            height=100
        )
    
    with col9:
        gia_tri_tai_san = st.number_input(
            "Giá trị định giá (VNĐ)",
            min_value=0,
            value=int(extracted_data['gia_tri_tai_san']),
            step=1000000,
            format="%d"
        )
        st.caption(f"💵 {format_number(gia_tri_tai_san)} VNĐ")
    
    st.markdown('<div class="sub-header">💼 Thông tin Thu nhập & Chi phí</div>', unsafe_allow_html=True)
    
    col10, col11 = st.columns(2)
    with col10:
        thu_nhap_thang = st.number_input(
            "Thu nhập hàng tháng (VNĐ)",
            min_value=0,
            value=int(extracted_data['thu_nhap_thang']),
            step=1000000,
            format="%d"
        )
        st.caption(f"💵 {format_number(thu_nhap_thang)} VNĐ")
    
    with col11:
        chi_phi_thang = st.number_input(
            "Chi phí hàng tháng (VNĐ)",
            min_value=0,
            value=int(extracted_data['chi_phi_thang']),
            step=1000000,
            format="%d"
        )
        st.caption(f"💵 {format_number(chi_phi_thang)} VNĐ")
    
    # Lưu dữ liệu vào session state
    st.session_state.customer_data = {
        'ho_ten': ho_ten,
        'cccd': cccd,
        'dia_chi': dia_chi,
        'sdt': sdt,
        'muc_dich_vay': muc_dich_vay,
        'tong_nhu_cau_von': tong_nhu_cau_von,
        'von_doi_ung': von_doi_ung,
        'so_tien_vay': so_tien_vay,
        'lai_suat': lai_suat,
        'thoi_gian_vay': thoi_gian_vay,
        'mo_ta_tai_san': mo_ta_tai_san,
        'gia_tri_tai_san': gia_tri_tai_san,
        'thu_nhap_thang': thu_nhap_thang,
        'chi_phi_thang': chi_phi_thang
    }

# Tab 2: Phân tích Chỉ số
with tab2:
    if 'customer_data' in st.session_state:
        data = st.session_state.customer_data
        
        st.markdown('<div class="sub-header">📊 Các Chỉ số Tài chính Quan trọng</div>', unsafe_allow_html=True)
        
        # Tính toán các chỉ số
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            ty_le_vay = (data['so_tien_vay'] / data['tong_nhu_cau_von'] * 100) if data['tong_nhu_cau_von'] > 0 else 0
            st.metric(
                "Tỷ lệ Vay/Tổng nhu cầu",
                f"{ty_le_vay:.2f}%"
            )
        
        with col2:
            ty_le_doi_ung = (data['von_doi_ung'] / data['tong_nhu_cau_von'] * 100) if data['tong_nhu_cau_von'] > 0 else 0
            st.metric(
                "Tỷ lệ Vốn đối ứng",
                f"{ty_le_doi_ung:.2f}%"
            )
        
        with col3:
            ltv = (data['so_tien_vay'] / data['gia_tri_tai_san'] * 100) if data['gia_tri_tai_san'] > 0 else 0
            st.metric(
                "LTV (Loan to Value)",
                f"{ltv:.2f}%"
            )
        
        with col4:
            thu_nhap_rong = data['thu_nhap_thang'] - data['chi_phi_thang']
            st.metric(
                "Thu nhập ròng/tháng",
                f"{format_number(thu_nhap_rong)} VNĐ"
            )
        
        st.markdown('<div class="sub-header">📅 Kế hoạch Trả nợ Chi tiết</div>', unsafe_allow_html=True)
        
        if data['so_tien_vay'] > 0 and data['thoi_gian_vay'] > 0:
            schedule_df = calculate_repayment_schedule(
                data['so_tien_vay'],
                data['lai_suat'],
                data['thoi_gian_vay']
            )
            
            if not schedule_df.empty:
                st.dataframe(
                    schedule_df,
                    use_container_width=True,
                    height=400
                )
                
                # Tổng hợp
                col1, col2, col3 = st.columns(3)
                
                # Parse số từ chuỗi đã format
                tong_goc = sum([parse_number(x) for x in schedule_df['Gốc trả']])
                tong_lai = sum([parse_number(x) for x in schedule_df['Lãi trả']])
                tong_thanh_toan = tong_goc + tong_lai
                
                with col1:
                    st.info(f"**Tổng gốc:** {format_number(tong_goc)} VNĐ")
                with col2:
                    st.info(f"**Tổng lãi:** {format_number(tong_lai)} VNĐ")
                with col3:
                    st.info(f"**Tổng thanh toán:** {format_number(tong_thanh_toan)} VNĐ")
        else:
            st.warning("⚠️ Vui lòng nhập đầy đủ thông tin vay để xem kế hoạch trả nợ!")
    else:
        st.info("ℹ️ Vui lòng nhập thông tin ở Tab 'Nhập liệu & Trích xuất' trước.")

# Tab 3: Biểu đồ Trực quan
with tab3:
    if 'customer_data' in st.session_state:
        data = st.session_state.customer_data
        
        st.markdown('<div class="sub-header">📊 Cơ cấu Nguồn vốn</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Biểu đồ tròn cơ cấu vốn
            if data['tong_nhu_cau_von'] > 0:
                fig_pie = go.Figure(data=[go.Pie(
                    labels=['Vốn vay', 'Vốn đối ứng'],
                    values=[data['so_tien_vay'], data['von_doi_ung']],
                    hole=0.3,
                    marker_colors=['#1f77b4', '#ff7f0e']
                )])
                fig_pie.update_layout(
                    title="Cơ cấu Nguồn vốn",
                    height=400
                )
                st.plotly_chart(fig_pie, use_container_width=True)
        
        with col2:
            # Biểu đồ thu nhập vs chi phí
            fig_bar = go.Figure(data=[
                go.Bar(name='Thu nhập', x=['Hàng tháng'], y=[data['thu_nhap_thang']], marker_color='#2ecc71'),
                go.Bar(name='Chi phí', x=['Hàng tháng'], y=[data['chi_phi_thang']], marker_color='#e74c3c'),
                go.Bar(name='Thu nhập ròng', x=['Hàng tháng'], y=[data['thu_nhap_thang'] - data['chi_phi_thang']], marker_color='#3498db')
            ])
            fig_bar.update_layout(
                title="Thu nhập & Chi phí",
                yaxis_title="VNĐ",
                barmode='group',
                height=400
            )
            st.plotly_chart(fig_bar, use_container_width=True)
        
        # Biểu đồ dư nợ giảm dần
        if data['so_tien_vay'] > 0 and data['thoi_gian_vay'] > 0:
            st.markdown('<div class="sub-header">📉 Biểu đồ Dư nợ Giảm dần</div>', unsafe_allow_html=True)
            
            schedule_df = calculate_repayment_schedule(
                data['so_tien_vay'],
                data['lai_suat'],
                data['thoi_gian_vay']
            )
            
            if not schedule_df.empty:
                # Convert string to number for plotting
                du_no_values = [parse_number(x) for x in schedule_df['Dư nợ cuối kỳ']]
                
                fig_line = go.Figure()
                fig_line.add_trace(go.Scatter(
                    x=schedule_df['Kỳ'],
                    y=du_no_values,
                    mode='lines+markers',
                    name='Dư nợ',
                    line=dict(color='#e74c3c', width=2),
                    marker=dict(size=6)
                ))
                fig_line.update_layout(
                    title="Dư nợ Giảm dần theo Thời gian",
                    xaxis_title="Kỳ trả nợ",
                    yaxis_title="Dư nợ (VNĐ)",
                    height=400,
                    hovermode='x unified'
                )
                st.plotly_chart(fig_line, use_container_width=True)
    else:
        st.info("ℹ️ Vui lòng nhập thông tin ở Tab 'Nhập liệu & Trích xuất' trước.")

# Tab 4: Phân tích bởi AI
with tab4:
    if not api_key:
        st.warning("⚠️ Vui lòng nhập Gemini API Key ở thanh bên để sử dụng tính năng AI!")
        st.info("💡 **Hướng dẫn lấy API Key:** Mở sidebar → Nhấn vào 'Hướng dẫn lấy Gemini API Key'")
    elif 'customer_data' not in st.session_state:
        st.info("ℹ️ Vui lòng nhập thông tin ở Tab 'Nhập liệu & Trích xuất' trước.")
    else:
        data = st.session_state.customer_data
        
        st.markdown('<div class="sub-header">🤖 Phân tích Thông minh bằng AI</div>', unsafe_allow_html=True)
        
        if st.button("🚀 Bắt đầu Phân tích", type="primary"):
            with st.spinner('🔄 Đang phân tích...'):
                try:
                    model = genai.GenerativeModel('gemini-2.0-flash-exp')
                    
                    # Phân tích 1: Dựa trên file gốc
                    st.markdown("### 📄 Phân tích 1 - Dựa trên File gốc")
                    st.caption("*Nguồn dữ liệu: Phân tích từ file .docx của khách hàng*")
                    
                    if st.session_state.uploaded_content:
                        prompt1 = f"""
Bạn là chuyên gia thẩm định tín dụng ngân hàng. Hãy phân tích phương án kinh doanh sau:

{st.session_state.uploaded_content}

Hãy đưa ra nhận định về:
1. Tổng quan về phương án
2. Điểm mạnh của phương án
3. Điểm yếu và rủi ro
4. Đề xuất cải thiện (nếu có)

Trả lời bằng tiếng Việt, ngắn gọn và chuyên nghiệp.
"""
                        response1 = model.generate_content(prompt1)
                        st.markdown(response1.text)
                    else:
                        st.warning("Chưa có file upload để phân tích.")
                    
                    st.markdown("---")
                    
                    # Phân tích 2: Dựa trên dữ liệu đã hiệu chỉnh
                    st.markdown("### 📊 Phân tích 2 - Dựa trên Dữ liệu đã Hiệu chỉnh")
                    st.caption("*Nguồn dữ liệu: Phân tích từ các thông số và chỉ số đã tính toán trên ứng dụng*")
                    
                    ty_le_vay = (data['so_tien_vay'] / data['tong_nhu_cau_von'] * 100) if data['tong_nhu_cau_von'] > 0 else 0
                    ty_le_doi_ung = (data['von_doi_ung'] / data['tong_nhu_cau_von'] * 100) if data['tong_nhu_cau_von'] > 0 else 0
                    ltv = (data['so_tien_vay'] / data['gia_tri_tai_san'] * 100) if data['gia_tri_tai_san'] > 0 else 0
                    thu_nhap_rong = data['thu_nhap_thang'] - data['chi_phi_thang']
                    
                    # Tính khoản trả hàng tháng
                    if data['so_tien_vay'] > 0 and data['thoi_gian_vay'] > 0:
                        lai_suat_thang = data['lai_suat'] / 12 / 100
                        goc_tra = data['so_tien_vay'] / data['thoi_gian_vay']
                        lai_tra_ky_dau = data['so_tien_vay'] * lai_suat_thang
                        tong_tra_ky_dau = goc_tra + lai_tra_ky_dau
                        
                        dscr = thu_nhap_rong / tong_tra_ky_dau if tong_tra_ky_dau > 0 else 0
                    else:
                        tong_tra_ky_dau = 0
                        dscr = 0
                    
                    prompt2 = f"""
Bạn là chuyên gia thẩm định tín dụng. Phân tích khoản vay với các thông số sau:

THÔNG TIN KHÁCH HÀNG:
- Họ tên: {data['ho_ten']}
- Mục đích vay: {data['muc_dich_vay']}

THÔNG TIN TÀI CHÍNH:
- Tổng nhu cầu vốn: {format_number(data['tong_nhu_cau_von'])} VNĐ
- Vốn đối ứng: {format_number(data['von_doi_ung'])} VNĐ ({ty_le_doi_ung:.2f}%)
- Số tiền vay: {format_number(data['so_tien_vay'])} VNĐ ({ty_le_vay:.2f}%)
- Lãi suất: {data['lai_suat']}%/năm
- Thời hạn: {data['thoi_gian_vay']} tháng

THU NHẬP & CHI PHÍ:
- Thu nhập tháng: {format_number(data['thu_nhap_thang'])} VNĐ
- Chi phí tháng: {format_number(data['chi_phi_thang'])} VNĐ
- Thu nhập ròng: {format_number(thu_nhap_rong)} VNĐ
- Khoản trả nợ kỳ đầu: {format_number(tong_tra_ky_dau)} VNĐ
- DSCR (Debt Service Coverage Ratio): {dscr:.2f}

TÀI SẢN ĐẢM BẢO:
- Mô tả: {data['mo_ta_tai_san']}
- Giá trị: {format_number(data['gia_tri_tai_san'])} VNĐ
- LTV: {ltv:.2f}%

Hãy đánh giá:
1. Tính khả thi tài chính của khoản vay
2. Khả năng trả nợ của khách hàng
3. Mức độ rủi ro và các yếu tố cần lưu ý
4. Kiến nghị chấp thuận/từ chối/điều chỉnh

Trả lời bằng tiếng Việt, chuyên nghiệp và chi tiết.
"""
                    response2 = model.generate_content(prompt2)
                    st.markdown(response2.text)
                    
                except Exception as e:
                    error_msg = str(e)
                    
                    if "API_KEY_INVALID" in error_msg or "expired" in error_msg.lower():
                        st.error("❌ **API Key không hợp lệ hoặc đã hết hạn!**")
                        st.markdown("""
                        <div class='api-note'>
                            <strong>💡 Giải pháp:</strong><br>
                            1. Mở sidebar (thanh bên trái)<br>
                            2. Tạo API key mới tại: <a href='https://aistudio.google.com/app/apikey' target='_blank'>Google AI Studio</a><br>
                            3. Copy API key mới và paste vào ô "Gemini API Key"<br>
                            4. Thử lại phân tích
                        </div>
                        """, unsafe_allow_html=True)
                    elif "quota" in error_msg.lower() or "rate" in error_msg.lower():
                        st.error("❌ **Đã vượt quá giới hạn sử dụng API!**")
                        st.info("💡 Vui lòng đợi 1 phút hoặc tạo API key mới.")
                    else:
                        st.error(f"❌ **Lỗi khi gọi API:**")
                        st.code(error_msg)
                        st.info("💡 Vui lòng kiểm tra lại API key hoặc kết nối Internet.")

# Tab 5: Chatbot
with tab5:
    if not api_key:
        st.warning("⚠️ Vui lòng nhập Gemini API Key ở thanh bên để sử dụng Chatbot!")
        st.info("💡 **Hướng dẫn lấy API Key:** Mở sidebar → Nhấn vào 'Hướng dẫn lấy Gemini API Key'")
    else:
        st.markdown('<div class="sub-header">💬 Chatbot Hỗ trợ Thẩm định</div>', unsafe_allow_html=True)
        
        # Nút xóa lịch sử
        col1, col2 = st.columns([6, 1])
        with col2:
            if st.button("🗑️ Xóa lịch sử"):
                st.session_state.chat_history = []
                st.rerun()
        
        # Hiển thị lịch sử chat
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Input chat
        if prompt := st.chat_input("Hỏi gì đó về phương án..."):
            # Thêm tin nhắn người dùng
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Tạo context từ dữ liệu hiện tại
            context = ""
            if 'customer_data' in st.session_state:
                data = st.session_state.customer_data
                context = f"""
THÔNG TIN PHƯƠNG ÁN HIỆN TẠI:
- Khách hàng: {data['ho_ten']}
- Mục đích: {data['muc_dich_vay']}
- Số tiền vay: {format_number(data['so_tien_vay'])} VNĐ
- Lãi suất: {data['lai_suat']}%/năm
- Thời hạn: {data['thoi_gian_vay']} tháng
- Thu nhập tháng: {format_number(data['thu_nhap_thang'])} VNĐ
- Chi phí tháng: {format_number(data['chi_phi_thang'])} VNĐ
"""
            
            # Gọi AI
            try:
                model = genai.GenerativeModel('gemini-2.0-flash-exp')
                
                full_prompt = f"""
Bạn là trợ lý AI chuyên về thẩm định tín dụng ngân hàng. 

{context}

Câu hỏi: {prompt}

Trả lời ngắn gọn, chuyên nghiệp bằng tiếng Việt.
"""
                
                response = model.generate_content(full_prompt)
                assistant_response = response.text
                
                # Thêm phản hồi vào lịch sử
                st.session_state.chat_history.append({"role": "assistant", "content": assistant_response})
                
                with st.chat_message("assistant"):
                    st.markdown(assistant_response)
                    
            except Exception as e:
                error_msg = str(e)
                
                if "API_KEY_INVALID" in error_msg or "expired" in error_msg.lower():
                    error_response = """❌ **API Key không hợp lệ hoặc đã hết hạn!**

💡 **Giải pháp:**
1. Mở sidebar (thanh bên trái)
2. Tạo API key mới tại: [Google AI Studio](https://aistudio.google.com/app/apikey)
3. Copy API key mới và paste vào ô "Gemini API Key"
4. Thử lại câu hỏi"""
                    
                elif "quota" in error_msg.lower() or "rate" in error_msg.lower():
                    error_response = "❌ Đã vượt quá giới hạn sử dụng API! Vui lòng đợi 1 phút hoặc tạo API key mới."
                else:
                    error_response = f"❌ Lỗi: {error_msg}\n\n💡 Vui lòng kiểm tra lại API key hoặc kết nối Internet."
                
                st.session_state.chat_history.append({"role": "assistant", "content": error_response})
                with st.chat_message("assistant"):
                    st.markdown(error_response)

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #666; padding: 1rem;'>
        <p>© 2024 Hệ thống Thẩm định Phương án Kinh doanh | Phát triển bởi AI</p>
    </div>
    """,
    unsafe_allow_html=True
)
